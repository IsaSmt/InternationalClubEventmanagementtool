# modules/report_ai_generator.py
import ollama
from docx import Document
from docx.shared import Pt
from io import BytesIO
import os
import time
import streamlit as st 

OLLAMA_MODEL_NAME = "phi3:mini"

# Diese Funktion prüft, ob das Ollama-Modell lokal verfügbar ist.
def _check_ollama_model_availability_for_cli(model_name_to_check: str) -> bool:
    try:
        models_list = ollama.list().get('models', [])
        if not models_list: return False
        for model_dict in models_list:
            name_from_list_raw = model_dict.get('name') or model_dict.get('model')
            if name_from_list_raw and name_from_list_raw.strip() == model_name_to_check.strip():
                return True
        return False
    except ConnectionRefusedError:
        print("CLI-FEHLER: Verbindung zu Ollama fehlgeschlagen. Läuft der Dienst?")
        return False
    except Exception as e:
        print(f"CLI-FEHLER bei Modellprüfung: {type(e).__name__} - {e}")
        return False

def _pull_ollama_model_for_cli(model_name: str, progress_placeholder_unused):
    try:
        current_status = ""
        print(f"CLI: Starte Download für Modell '{model_name}'...")
        for progress in ollama.pull(model_name, stream=True):
            status = progress.get("status", "")
            if status != current_status:
                current_status = status
                print(f"CLI Ollama Pull Status: {status}")
        print(f"CLI: Modell '{model_name}' Download abgeschlossen/verifiziert.")
        return True
    except Exception as e:
        print(f"CLI-FEHLER beim Download des Modells '{model_name}': {type(e).__name__} - {e}")
        return False

def check_ollama_model_availability(model_name_to_check: str) -> bool:
    try:
        models_list = ollama.list().get('models', [])
        if not models_list: return False
        for model_dict in models_list:
            name_from_list_raw = model_dict.get('name') or model_dict.get('model')
            if name_from_list_raw and name_from_list_raw.strip() == model_name_to_check.strip():
                return True
        return False
    except ConnectionRefusedError:
        st.error("Verbindung zum Ollama-Dienst fehlgeschlagen.")
        return False
    except Exception as e:
        st.error(f"Fehler bei Modellprüfung mit Ollama: {e}")
        return False

def pull_ollama_model(model_name: str, progress_placeholder):
    try:
        current_status = ""
        for progress in ollama.pull(model_name, stream=True):
            status = progress.get("status", "")
            if status != current_status and progress_placeholder:
                current_status = status
                progress_placeholder.info(f"Lade Modell '{model_name}': {status}")
        if progress_placeholder:
            progress_placeholder.success(f"Modell '{model_name}' erfolgreich heruntergeladen/verifiziert!")
            time.sleep(2)
            progress_placeholder.empty()
        return True
    except ollama.ResponseError as e:
        if progress_placeholder: progress_placeholder.empty()
        st.error(f"Fehler von der Ollama API beim Pull-Versuch für '{model_name}': {e}")
        return False
    except Exception as e:
        if progress_placeholder: progress_placeholder.empty()
        st.error(f"Allgemeiner Fehler beim Herunterladen des Modells '{model_name}': {type(e).__name__} - {e}")
        return False

def streamlit_is_running():
    try:
        return hasattr(st, 'secrets') and callable(st.secrets.get)
    except Exception:
        return False

def generate_experience_report_docx(tutor_freitext: str, event_title: str) -> bytes | None:
    """
    Generiert einen Erfahrungsbericht als Word-Datei (.docx) basierend auf dem Freitext
    unter Verwendung eines lokalen LLMs über Ollama.
    """
    is_streamlit_context = streamlit_is_running()
    model_status_placeholder = None

    if is_streamlit_context:
        model_status_placeholder = st.empty()
        if not tutor_freitext:
            model_status_placeholder.warning("Freitext für den Bericht ist leer.")
            return None
    else: 
        if not tutor_freitext:
            print("WARNUNG (CLI): Freitext für den Bericht ist leer.")
            return None
    
    try:
        ollama.list() 
    except ConnectionRefusedError:
        msg = "Verbindung zum Ollama-Dienst fehlgeschlagen. Bitte Ollama starten."
        if is_streamlit_context and model_status_placeholder: model_status_placeholder.error(msg)
        else: print(f"FEHLER (CLI): {msg}")
        return None
    except Exception as e:
        msg = f"Fehler bei der Kommunikation mit Ollama: {e}"
        if is_streamlit_context and model_status_placeholder: model_status_placeholder.error(msg)
        else: print(f"FEHLER (CLI): {msg}")
        return None

    model_available_func = check_ollama_model_availability if is_streamlit_context else _check_ollama_model_availability_for_cli
    pull_model_func = pull_ollama_model if is_streamlit_context else _pull_ollama_model_for_cli
    
    if not model_available_func(OLLAMA_MODEL_NAME):
        msg_info = f"Modell '{OLLAMA_MODEL_NAME}' nicht lokal. Starte Download..."
        if is_streamlit_context and model_status_placeholder: model_status_placeholder.info(msg_info)
        else: print(f"INFO (CLI): {msg_info}")
        
        pull_progress_arg = model_status_placeholder if is_streamlit_context else None
        if not pull_model_func(OLLAMA_MODEL_NAME, pull_progress_arg): 
            return None 
        if not model_available_func(OLLAMA_MODEL_NAME): 
            msg_error = f"Modell '{OLLAMA_MODEL_NAME}' auch nach Download nicht gefunden."
            if is_streamlit_context and model_status_placeholder: model_status_placeholder.error(msg_error)
            else: print(f"FEHLER (CLI): {msg_error}")
            return None
            
    if is_streamlit_context and model_status_placeholder: model_status_placeholder.empty() 

    prompt = f"""
    Erstelle einen kurzen, prägnanten und professionellen Erfahrungsbericht von etwa einer halben Seite (ca. 150-250 Wörter)
    im .docx-Format basierend auf den folgenden Stichpunkten und Informationen.
    Der Bericht ist für interne Zwecke und soll einen guten Überblick über das Event geben.
    Stil: Sachlich, positiv (wenn möglich), aber auch ehrliche Nennung von Problemen, falls vorhanden.
    Struktur: Kurze Einleitung, Hauptteil (Ablauf, Highlights, ggf. Probleme), kurzes Fazit/Ausblick.
    Event-Titel: {event_title}
    Stichpunkte/Freitext:
    ---
    {tutor_freitext}
    ---
    Bitte generiere nur den reinen Text für den Bericht, ohne zusätzliche Anmerkungen wie "Hier ist der Bericht:" etc.
    Beginne direkt mit dem Berichtstext. Achte auf eine klare Absatzstruktur.
    """
    try:
        ai_text_response = None
        spinner_text = f"Bericht wird mit KI-Modell '{OLLAMA_MODEL_NAME}' generiert..."
        if is_streamlit_context:
            with st.spinner(spinner_text):
                response = ollama.chat(model=OLLAMA_MODEL_NAME, messages=[{'role': 'user', 'content': prompt}])
                ai_text_response = response['message']['content']
        else: 
            print(f"INFO (CLI): {spinner_text}")
            response = ollama.chat(model=OLLAMA_MODEL_NAME, messages=[{'role': 'user', 'content': prompt}])
            ai_text_response = response['message']['content']

        if not ai_text_response:
            raise ValueError("KI hat keinen Text zurückgegeben.")

        document = Document()
        document.add_heading(f'Erfahrungsbericht: {event_title}', level=1)
        style = document.styles['Normal']
        font = style.font
        try: 
            font.name = 'Nunito' 
        except: 
            font.name = 'Calibri'
            if not is_streamlit_context: print("WARNUNG (Bericht-CLI): Nunito nicht für DOCX gefunden, verwende Calibri.")
        font.size = Pt(11)
        
        paragraphs = ai_text_response.strip().split('\n\n') 
        if len(paragraphs) == 1 and '\n' in ai_text_response: 
            paragraphs = ai_text_response.strip().split('\n')
        for para_text in paragraphs:
            if para_text.strip(): 
                document.add_paragraph(para_text.strip(), style='Normal')

        bio = BytesIO()
        document.save(bio)
        bio.seek(0)
        return bio.getvalue()

    except Exception as e: 
        msg_error = f"Fehler bei der KI-Berichtsgenerierung oder Docx-Erstellung: {type(e).__name__} - {e}"
        if is_streamlit_context: st.error(msg_error)
        else: print(f"FEHLER (CLI): {msg_error}")
        return None

if __name__ == "__main__":
    print("Starte Testlauf für modules/report_ai_generator.py...")
    print(f"Verwendet Ollama Modell: {OLLAMA_MODEL_NAME}")
    test_event_title_ai = "Canyoning International Club SS25"
    test_tutor_freitext_ai = """
    Um 9 Uhr waren wir in München am Treffpunkt, sind dann losgefahren zum Canyoning Startpunkt. 
    1.5 Std. später waren wir da und haben uns Neoprenanzüge angezogen. 
    Dann sind wir 40 Minuten gewandert zum Startpunkt in den Canyon. 
    Dann haben wir uns an zwei verschiedenen hohen Wänden mit 20-30 Metern abgeseilt 
    und sind 2 km durch den Canyon gestiegen. Das Wetter war sehr heiß und sonnig, 
    aber man hatte durch das kühle Wasser gute Abkühlung. 
    Gegen 15 Uhr sind wir wieder zurück nach München gefahren. 
    """
    print("\nVersuche, Erfahrungsbericht zu generieren (CLI-Modus)...")
    try:
        report_bytes = generate_experience_report_docx(test_tutor_freitext_ai, test_event_title_ai)
        if report_bytes:
            output_dir = "output"
            if not os.path.exists(output_dir): os.makedirs(output_dir)
            safe_event_title = "".join(c if c.isalnum() else "_" for c in test_event_title_ai)
            report_test_filename = os.path.join(output_dir, f"TEST_CLI_Erfahrungsbericht_{safe_event_title}.docx")
            with open(report_test_filename, "wb") as f: f.write(report_bytes)
            print(f"\nTEST ERFOLGREICH: Bericht als '{report_test_filename}' gespeichert.")
        else: print("\nFEHLER im Test: KI-Bericht konnte nicht generiert werden.")
    except Exception as e_global: print(f"\nFEHLER im Test: {type(e_global).__name__} - {e_global}")
    print("\nTestlauf für modules/report_ai_generator.py beendet.")